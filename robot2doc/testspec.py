#!/env/python

import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK_TYPE

import config as cfg

DOC_CLAUSE_LVL_1 = cfg.DOC_CLAUSE_LVL_1
DOC_CLAUSE_LVL_2 = cfg.DOC_CLAUSE_LVL_2

ANNEX_TITLE_LEVEL = cfg.ANNEX_TITLE_LEVEL
DOC_SUITE_LEVEL = cfg.DOC_SUITE_LEVEL
DOC_TC_LEVEL = cfg.DOC_TC_LEVEL

DOC_TC_FONT_SIZE = cfg.DOC_TC_FONT_SIZE
DOC_TC_FONT_NAME = cfg.DOC_TC_FONT_NAME

TP_TITLE = 'Test Purpose'
TC_TITLE = 'Test Case'

TABLE_STYLE = 'TestTable'

class TestSpec():

    def __init__(self, path=None):
        if path:
            print("Opening doc: "+path)
        print("Current dir: "+os.getcwd())
        self.doc = Document(path)

    @staticmethod
    def cell_text_bold(cell):
        cell.paragraphs[0].runs[0].font.bold = True

    @staticmethod
    def cell_text_centered(cell):
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    @staticmethod
    def cell_text_mono(cell):
        cell.paragraphs[0].runs[0].font.name = DOC_TC_FONT_NAME
        cell.paragraphs[0].runs[0].font.size = Pt(DOC_TC_FONT_SIZE)

    @staticmethod
    def mk_ref(lv1, lv2="", lv3="", lv4="", lv5=""):
        secs = [str(x) for x in [lv1, lv2, lv3, lv4, lv5] if x != ""]
        return (".".join(secs))

    @staticmethod
    def mk_heading(txt, sec, subsec1="", subsec2="", subsec3=""):
        return TestSpec.mk_ref(sec, subsec1, subsec2, subsec3) + " " + txt

    @staticmethod
    def mk_tp_hdr(table):
        hdr_cells = table.rows[0].cells
        hdr_cells[0].merge(hdr_cells[1])
        hdr_cells[0].text = TP_TITLE
        TestSpec.cell_text_bold(hdr_cells[0])
        TestSpec.cell_text_centered(hdr_cells[0])

    @staticmethod
    def mk_tc_hdr(table):
        hdr_cells = table.add_row().cells
        hdr_cells[0].merge(hdr_cells[1])
        hdr_cells[0].text = TC_TITLE
        TestSpec.cell_text_bold(hdr_cells[0])
        TestSpec.cell_text_centered(hdr_cells[0])

    def add_commit_url(self, commit : str, robot_file : str):
        '''
        Adds a note to the document containing the URL to the location of the
        file, according to the configured URL prefix configured.
        '''
        commit_string="https://forge.etsi.org/gitlab/nfv/api-tests/raw/"+commit+"/SOL00Y/AAA-API/"+robot_file
        self.doc.add_paragraph("Note: Robot code can be found at " + commit_string)

    def add_tp(self, fields, testbehaviour):
        table = self.doc.add_table(cols=2, rows=1)
        table.style = TABLE_STYLE
        table.autofit = True
        TestSpec.mk_tp_hdr(table)
        for tpf in fields:
            row_cells = table.add_row().cells
            row_cells[0].width = 2
            row_cells[0].text = tpf.key
            TestSpec.cell_text_bold(row_cells[0])
            row_cells[1].text = tpf.value
        TestSpec.mk_tc_hdr(table)
        row_cells = table.add_row().cells
        row_cells[0].merge(row_cells[1])
        row_cells[0].text = testbehaviour
        TestSpec.cell_text_mono(row_cells[0])

    def add_heading(self, txt, lvl, sec="1", subsec1="", subsec2="", subsec3=""):
        self.doc.add_heading(TestSpec.mk_heading(txt, sec, subsec1, subsec2, subsec3), lvl)

    def save(self, path):
        self.doc.save(path)

    def add_main_heading(self, txt):
        self.doc.add_page_break()
        self.doc.add_heading(txt, ANNEX_TITLE_LEVEL)

    def add_sub_heading(self, txt, sec, subsec1, subsec2, subsec3):
        '''
        Add a sub heading to the document
        @txt Title of the heading
        @sec Numer of the section of the heading
        '''
        self.doc.add_heading(TestSpec.mk_heading(txt, sec, subsec1, subsec2, subsec3), DOC_SUITE_LEVEL)
  
