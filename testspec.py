#!/env/python

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK_TYPE

ANNEX_LETTER = "B"
DOC_BASE_SEC = 1

ANNEX_TITLE_LEVEL = 1
DOC_SUITE_LEVEL = 2
DOC_TC_LEVEL    = 3

DOC_TC_FONT_SIZE = 8
DOC_TC_FONT_NAME = "Courier New"

TP_TITLE = 'Test Purpose'
TC_TITLE = 'Test Case'

TABLE_STYLE = 'TestTable'

class TestSpec():

    def __init__(self, path=None):
        self.doc = Document(path)

    @staticmethod
    def cell_text_bold(cell):
        cell.paragraphs[0].runs[0].font.bold = True

    @staticmethod
    def cell_text_centered(cell):
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    @staticmethod
    def cell_text_mono(cell):
        cell.paragraphs[0].runs[0].font.name = DOC_TC_FONT_NAME
        cell.paragraphs[0].runs[0].font.size = Pt(DOC_TC_FONT_SIZE)

    @staticmethod
    def mk_ref(sec, subsec=""):
        if subsec != "": 
            return ANNEX_LETTER +"."+ str(sec)+"."+str(subsec)
        return ANNEX_LETTER +"."+ str(sec)

    @staticmethod
    def mk_heading(txt, sec, subsec=""):
        return TestSpec.mk_ref(sec, subsec) + " " + txt

    @staticmethod
    def mk_tp_hdr(table):
        hdr_cells = table.rows[0].cells
        hdr_cells[0].merge(hdr_cells[1])
        hdr_cells[0].text = TP_TITLE
        TestSpec.cell_text_bold(hdr_cells[0])
        TestSpec.cell_text_centered(hdr_cells[0])

    @staticmethod
    def mk_tc_hdr(table):
        hdr_cells = table.add_row().cells
        hdr_cells[0].merge(hdr_cells[1])
        hdr_cells[0].text = TC_TITLE
        TestSpec.cell_text_bold(hdr_cells[0])
        TestSpec.cell_text_centered(hdr_cells[0])

    def add_tp(self, fields, testbehaviour):
        table = self.doc.add_table(cols=2, rows=1)
        table.style = TABLE_STYLE
        table.autofit = True
        TestSpec.mk_tp_hdr(table)
        for tpf in fields:
            row_cells = table.add_row().cells
            row_cells[0].width = 2
            row_cells[0].text = tpf.key
            TestSpec.cell_text_bold(row_cells[0])
            row_cells[1].text = tpf.value
        TestSpec.mk_tc_hdr(table)
        row_cells = table.add_row().cells
        row_cells[0].merge(row_cells[1])
        row_cells[0].text = testbehaviour
        TestSpec.cell_text_mono(row_cells[0])

    def add_heading(self, txt, lvl, sec="1", subsec=""):
        self.doc.add_heading(TestSpec.mk_heading(txt, sec, subsec), lvl)
        
    def save(self, path):
        self.doc.save(path)

    def add_main_heading(self, txt):
        self.doc.add_page_break()
        self.doc.add_heading(txt, ANNEX_TITLE_LEVEL)

    def add_sub_heading(self, txt, sec):
        '''
        Add a sub heading to the document
        @txt Title of the heading
        @sec Numer of the section of the heading
        '''
        self.doc.add_heading(TestSpec.mk_heading(txt,sec), DOC_SUITE_LEVEL)
        